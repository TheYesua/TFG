"""Servicio de exportación de Situaciones de Aprendizaje (CU-06).

Genera PDF (vía WeasyPrint) y DOCX (vía python-docx) a partir del
contenido JSON ya generado por la IA. La fuente de verdad es
``SituacionAprendizaje.contenido`` y el orden de las secciones lo dicta
``app.prompts.ORDEN_SECCIONES``.

La exportación NO regenera ni mutila contenido — sólo lo formatea.
"""
from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from flask import render_template
from weasyprint import HTML

from ..models import SituacionAprendizaje, Usuario
from ..prompts import ORDEN_SECCIONES


# Etiquetas humanas alineadas con SECCIONES de detalle.html
_ETIQUETAS = {
    "descripcion":         "Descripción de la situación",
    "objetivos":           "Objetivos didácticos",
    "conexion_curricular": "Conexión curricular",
    "secuencia_sesiones":  "Secuencia de sesiones",
    "evaluacion":          "Evaluación",
    "atencion_diversidad": "Atención a la diversidad",
}


def secciones_para_export() -> list[tuple[str, str]]:
    """Pares (clave, etiqueta) en el orden canónico LOMLOE."""
    return [(c, _ETIQUETAS.get(c, c.replace("_", " ").title())) for c in ORDEN_SECCIONES]


def renderizar_pdf(sa: SituacionAprendizaje, usuario: Usuario) -> bytes:
    """Devuelve los bytes del PDF generado para ``sa``.

    Reutiliza ``templates/exportacion/pdf.html``. WeasyPrint resuelve los
    estilos inline del propio template (no necesita acceso a /static).
    """
    html_str = render_template(
        "exportacion/pdf.html",
        sa=sa,
        usuario=usuario,
        secciones=secciones_para_export(),
        generado_en=datetime.now(timezone.utc).astimezone(),
    )
    # ``base_url`` permitiría a WeasyPrint resolver assets relativos si
    # se añadieran (logos, imágenes embebidas). De momento todo es CSS
    # inline, así que con None vale.
    pdf = HTML(string=html_str, base_url=None).write_pdf()
    if pdf is None:  # pragma: no cover - WeasyPrint nunca devuelve None
        raise RuntimeError("WeasyPrint no devolvió contenido PDF")
    return pdf


# ---------------------------------------------------------------------------
# DOCX (python-docx) — render programático
# ---------------------------------------------------------------------------

_COLOR_TITULO   = RGBColor(0x0F, 0x17, 0x2A)
_COLOR_LABEL    = RGBColor(0x37, 0x41, 0x51)
_COLOR_SUTIL    = RGBColor(0x6B, 0x72, 0x80)
_COLOR_PRIMARIO = RGBColor(0x5E, 0xD3, 0xD1)

# Fuente del DOCX. Inter (la usada en la web) no está instalada por defecto
# en Word, así que escogemos Calibri: es la fuente predeterminada de Office
# moderno, humanista y muy próxima visualmente a Inter. Sin esta declaración
# explícita, Word/LibreOffice caen al tema del template y han llegado a
# renderizar Cambria (serif), rompiendo la coherencia visual con la web.
_FONT_FAMILY = "DejaVu Sans"


def _force_font(element, name: str) -> None:
    """Fija ``rFonts`` en un elemento OOXML (estilo o run).

    python-docx expone ``font.name`` pero internamente sólo establece el
    atributo ``w:ascii``. Para que Word use SIEMPRE la misma fuente
    (occidental, complex script y eastAsia) hay que escribir las cuatro
    variantes; si no, en algunos sistemas se cae al tema del documento.
    """
    rpr = element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def _configurar_fuente_documento(doc: Document) -> None:
    """Aplica la familia tipográfica a los estilos base del documento.

    Cambia el estilo ``Normal`` y los de lista; los headings los seguimos
    fijando en cada ``_add_run`` por si el template trae overrides.
    """
    for style_name in ("Normal", "List Bullet", "List Number", "List Paragraph"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = _FONT_FAMILY
        _force_font(style.element, _FONT_FAMILY)


def _add_run(p, texto: str, *, bold=False, italic=False, size=None, color=None):
    """Helper: añade un run con propiedades comunes (y fuente fijada)."""
    r = p.add_run(texto)
    r.bold = bold
    r.italic = italic
    r.font.name = _FONT_FAMILY
    _force_font(r._element, _FONT_FAMILY)
    if size is not None:
        r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return r


def _add_h2(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    _add_run(p, texto, bold=True, size=15, color=_COLOR_TITULO)


def _add_h3(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    _add_run(p, texto, bold=True, size=11, color=_COLOR_LABEL)


def _add_label_value(doc: Document, etiqueta: str, valor: str) -> None:
    """Añade un párrafo con etiqueta destacada y valor a continuación."""
    if not valor:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _add_run(p, f"{etiqueta}: ", bold=True, size=10, color=_COLOR_LABEL)
    _add_run(p, valor, size=10)


def _add_lista(doc: Document, items, *, numerada=False) -> None:
    estilo = "List Number" if numerada else "List Bullet"
    for it in items or []:
        if not it:
            continue
        texto = it if isinstance(it, str) else str(it)
        p = doc.add_paragraph(texto, style=estilo)
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(10)
            r.font.name = _FONT_FAMILY
            _force_font(r._element, _FONT_FAMILY)


def _add_tabla(doc: Document, filas: list[list[str]], *, cabecera: list[str] | None = None) -> None:
    """Crea una tabla básica con bordes (estilo 'Table Grid')."""
    if not filas:
        return
    n_cols = max(len(cabecera or []), max(len(f) for f in filas))
    tabla = doc.add_table(rows=0, cols=n_cols)
    tabla.style = "Table Grid"
    tabla.autofit = True

    if cabecera:
        fila = tabla.add_row().cells
        for i, txt in enumerate(cabecera):
            fila[i].paragraphs[0].clear()
            _add_run(fila[i].paragraphs[0], txt, bold=True, size=9, color=_COLOR_LABEL)

    for f in filas:
        fila = tabla.add_row().cells
        for i, txt in enumerate(f):
            fila[i].paragraphs[0].clear()
            _add_run(fila[i].paragraphs[0], str(txt) if txt is not None else "—", size=9)

    # Espacio tras la tabla
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ----- secciones LOMLOE (en DOCX) -----

def _docx_descripcion(doc, d):
    if d.get("texto"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _add_run(p, d["texto"], size=10)
    if d.get("pregunta_guia"):
        _add_label_value(doc, "Pregunta guía", d["pregunta_guia"])
    if d.get("producto_final"):
        _add_label_value(doc, "Producto final", d["producto_final"])


def _docx_objetivos(doc, d):
    objetivos = d.get("objetivos") or []
    for i, o in enumerate(objetivos, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _add_run(p, f"{i}. ", bold=True, size=10, color=_COLOR_PRIMARIO)
        _add_run(p, o.get("texto", ""), size=10)
        if o.get("competencias"):
            comps = ", ".join(o["competencias"])
            _add_run(p, f"  [{comps}]", italic=True, size=9, color=_COLOR_SUTIL)


def _docx_conexion_curricular(doc, d):
    if d.get("competencias"):
        _add_h3(doc, "Competencias específicas")
        _add_tabla(
            doc,
            [[c.get("codigo", "—"), c.get("justificacion", "—")] for c in d["competencias"]],
            cabecera=["Código", "Justificación"],
        )
    if d.get("criterios"):
        _add_h3(doc, "Criterios de evaluación")
        _add_tabla(
            doc,
            [
                [c.get("codigo", "—"), c.get("competencia", "—"), c.get("justificacion", "—")]
                for c in d["criterios"]
            ],
            cabecera=["Código", "Comp.", "Justificación"],
        )
    if d.get("saberes"):
        _add_h3(doc, "Saberes básicos")
        _add_tabla(
            doc,
            [[s.get("codigo", "—"), s.get("justificacion", "—")] for s in d["saberes"]],
            cabecera=["Código", "Justificación"],
        )


def _docx_secuencia_sesiones(doc, d):
    for s in d.get("sesiones") or []:
        # Cabecera de sesión
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        _add_run(p, f"Sesión {s.get('numero', '?')}", bold=True, size=11, color=_COLOR_TITULO)
        if s.get("fase"):
            _add_run(p, f"  · {s['fase']}", italic=True, size=10, color=_COLOR_SUTIL)
        if s.get("duracion_minutos"):
            _add_run(p, f"  · {s['duracion_minutos']} min", size=10, color=_COLOR_SUTIL)
        if s.get("titulo"):
            ptit = doc.add_paragraph()
            ptit.paragraph_format.space_after = Pt(4)
            ptit.paragraph_format.keep_with_next = True
            _add_run(ptit, s["titulo"], italic=True, size=10)

        if s.get("actividades"):
            _add_label_value(doc, "Actividades", "")
            actividades = [
                a if isinstance(a, str) else a.get("descripcion", "")
                for a in s["actividades"]
            ]
            _add_lista(doc, [a for a in actividades if a], numerada=True)
        if s.get("recursos"):
            _add_label_value(doc, "Recursos", "")
            _add_lista(doc, s["recursos"])
        if s.get("criterios"):
            _add_label_value(doc, "Criterios trabajados", ", ".join(map(str, s["criterios"])))


def _docx_evaluacion(doc, d):
    if d.get("instrumentos"):
        _add_h3(doc, "Instrumentos de evaluación")
        _add_tabla(
            doc,
            [
                [
                    i.get("nombre", "—"),
                    i.get("momento", "—"),
                    f"{i['peso']}%" if i.get("peso") is not None else "—",
                ]
                for i in d["instrumentos"]
            ],
            cabecera=["Instrumento", "Momento", "Peso"],
        )
    if d.get("rubricas"):
        _add_h3(doc, "Rúbricas por criterio")
        NIVELES = [("excelente", "Excelente"), ("logrado", "Logrado"),
                   ("en_proceso", "En proceso"), ("iniciado", "Iniciado")]
        for r in d["rubricas"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            _add_run(p, f"Criterio {r.get('criterio', '—')}", bold=True, size=10, color=_COLOR_LABEL)
            niveles = r.get("niveles") or {}
            _add_tabla(
                doc,
                [[lab, niveles.get(k) or "—"] for k, lab in NIVELES],
            )


def _docx_atencion_diversidad(doc, d, tipo_adapt):
    if d.get("principios_dua"):
        _add_h3(doc, "Principios DUA aplicados")
        _add_lista(doc, d["principios_dua"])
    if d.get("ajustes_no_significativos") and tipo_adapt != "significativa":
        _add_h3(doc, "Ajustes no significativos")
        _add_lista(doc, d["ajustes_no_significativos"])
    if d.get("ajustes_significativos") and tipo_adapt != "no_significativa":
        _add_h3(doc, "Ajustes significativos")
        _add_lista(doc, d["ajustes_significativos"])


_RENDERS_DOCX = {
    "descripcion":         _docx_descripcion,
    "objetivos":           _docx_objetivos,
    "conexion_curricular": _docx_conexion_curricular,
    "secuencia_sesiones":  _docx_secuencia_sesiones,
    "evaluacion":          _docx_evaluacion,
}


def renderizar_docx(sa: SituacionAprendizaje, usuario: Usuario) -> bytes:
    """Devuelve los bytes del DOCX generado para ``sa``."""
    doc = Document()

    # Tipografía global ANTES de crear contenido. ``_add_run`` la repite por
    # run para garantizar coherencia incluso si el template tiene overrides.
    _configurar_fuente_documento(doc)

    # Márgenes A4 razonables
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.2)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    # Cabecera del documento
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run(p, sa.titulo, bold=True, size=20, color=_COLOR_TITULO)

    if sa.id_situacion_origen:
        etiqueta = (
            "ACS — Adaptación significativa"
            if sa.tipo_adaptacion == "significativa"
            else "ACNS — Adaptación no significativa"
        )
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _add_run(p, etiqueta, bold=True, size=9, color=RGBColor(0x92, 0x40, 0x0E))

    metas = []
    if sa.materia: metas.append(("Materia", sa.materia))
    if sa.curso: metas.append(("Curso", sa.curso))
    if sa.num_sesiones: metas.append(("Sesiones", str(sa.num_sesiones)))
    if sa.duracion_sesion_minutos:
        metas.append(("Duración", f"{sa.duracion_sesion_minutos} min"))
    metas.append(("Docente", usuario.nombre or "—"))
    if usuario.centro_educativo:
        metas.append(("Centro", usuario.centro_educativo))
    metas.append(("Generado", datetime.now(timezone.utc).astimezone().strftime("%d/%m/%Y %H:%M")))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    for i, (k, v) in enumerate(metas):
        if i:
            _add_run(p, "  ·  ", size=9, color=_COLOR_SUTIL)
        _add_run(p, f"{k}: ", bold=True, size=9, color=_COLOR_LABEL)
        _add_run(p, v, size=9)

    # Línea separadora
    sep = doc.add_paragraph("_" * 60)
    sep.runs[0].font.color.rgb = RGBColor(0xE5, 0xE7, 0xEB)
    sep.runs[0].font.size = Pt(8)
    sep.runs[0].font.name = _FONT_FAMILY
    _force_font(sep.runs[0]._element, _FONT_FAMILY)

    # Secciones
    contenido = sa.contenido or {}
    for clave, etiqueta in secciones_para_export():
        datos = contenido.get(clave)
        if not datos:
            continue
        _add_h2(doc, etiqueta)
        if clave == "atencion_diversidad":
            _docx_atencion_diversidad(doc, datos, sa.tipo_adaptacion)
        else:
            render = _RENDERS_DOCX.get(clave)
            if render:
                render(doc, datos)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def nombre_fichero(sa: SituacionAprendizaje, ext: str) -> str:
    """Nombre seguro para descarga: ``SA-{id}-{slug}.{ext}``."""
    slug = "".join(c if c.isalnum() else "-" for c in (sa.titulo or "").lower())
    slug = "-".join(filter(None, slug.split("-")))[:50] or "situacion"
    return f"SA-{sa.id_situacion}-{slug}.{ext}"
